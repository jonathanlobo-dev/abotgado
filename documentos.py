"""
aBOTgado - Motor de generación de documentos
=============================================
- Define las plantillas disponibles con sus campos
- Genera documentos .docx reemplazando marcadores
- Maneja el flujo de conversación (estado por usuario)
"""

import os
import re
import copy
import logging
from datetime import datetime
from docx import Document

import config

logger = logging.getLogger(__name__)

PLANTILLAS_DIR = config.PLANTILLAS_DIR


# ─── DEFINICIÓN DE PLANTILLAS ─────────────────────────────────────────────────
# Cada plantilla tiene:
#   nombre: nombre para mostrar al usuario
#   archivo: nombre del .docx en /plantillas/
#   campos: lista ordenada de campos a pedir
#   cada campo tiene: id, pregunta (lo que le preguntamos al usuario), ejemplo

PLANTILLAS = {
    "1": {
        "nombre": "Contrato de Arrendamiento de Vivienda",
        "archivo": "contrato_arrendamiento.docx",
        "descripcion": "Contrato entre arrendador e inquilino para alquiler de vivienda",
        "campos": [
            {"id": "ARRENDADOR_NOMBRE",  "pregunta": "Nombre completo del ARRENDADOR (dueño del inmueble):", "ejemplo": "María López Rodríguez"},
            {"id": "ARRENDADOR_CEDULA",  "pregunta": "Cédula del arrendador:", "ejemplo": "V-12.345.678"},
            {"id": "ARRENDATARIO_NOMBRE","pregunta": "Nombre completo del ARRENDATARIO (inquilino):", "ejemplo": "Carlos Pérez García"},
            {"id": "ARRENDATARIO_CEDULA","pregunta": "Cédula del arrendatario:", "ejemplo": "V-23.456.789"},
            {"id": "DIRECCION_INMUEBLE", "pregunta": "Dirección completa del inmueble:", "ejemplo": "Av. Principal, Edificio Sol, Piso 3, Apto 3-A, Urbanización Las Mercedes"},
            {"id": "CIUDAD",             "pregunta": "Ciudad:", "ejemplo": "Caracas"},
            {"id": "CANON_MENSUAL",      "pregunta": "Canon de arrendamiento mensual (monto):", "ejemplo": "200 dólares americanos"},
            {"id": "CANON_LETRAS",       "pregunta": "Canon en letras:", "ejemplo": "doscientos dólares americanos"},
            {"id": "DURACION_MESES",     "pregunta": "Duración del contrato en meses:", "ejemplo": "12"},
            {"id": "FECHA_INICIO",       "pregunta": "Fecha de inicio del contrato:", "ejemplo": "01 de abril de 2026"},
            {"id": "FECHA_FIRMA",        "pregunta": "Fecha de firma del documento:", "ejemplo": "17 de marzo de 2026"},
        ]
    },
    "2": {
        "nombre": "Carta de Renuncia Voluntaria",
        "archivo": "carta_renuncia.docx",
        "descripcion": "Carta formal de renuncia a un empleo con solicitud de prestaciones",
        "campos": [
            {"id": "TRABAJADOR_NOMBRE",  "pregunta": "Tu nombre completo:", "ejemplo": "José Hernández Díaz"},
            {"id": "TRABAJADOR_CEDULA",  "pregunta": "Tu cédula:", "ejemplo": "V-15.678.901"},
            {"id": "EMPRESA_NOMBRE",     "pregunta": "Nombre de la empresa:", "ejemplo": "Inversiones ABC, C.A."},
            {"id": "CARGO",              "pregunta": "Tu cargo actual:", "ejemplo": "Analista de Sistemas"},
            {"id": "FECHA_INGRESO",      "pregunta": "Fecha en que ingresaste a la empresa:", "ejemplo": "15 de enero de 2023"},
            {"id": "FECHA_RENUNCIA",     "pregunta": "Fecha efectiva de la renuncia:", "ejemplo": "30 de abril de 2026"},
            {"id": "CIUDAD",             "pregunta": "Ciudad:", "ejemplo": "Valencia"},
            {"id": "FECHA_FIRMA",        "pregunta": "Fecha de hoy:", "ejemplo": "17 de marzo de 2026"},
        ]
    },
    "3": {
        "nombre": "Poder Notarial General",
        "archivo": "poder_notarial.docx",
        "descripcion": "Poder para que otra persona actúe en tu nombre ante instituciones",
        "campos": [
            {"id": "PODERDANTE_NOMBRE",  "pregunta": "Tu nombre completo (quien otorga el poder):", "ejemplo": "Ana María Gutiérrez"},
            {"id": "PODERDANTE_CEDULA",  "pregunta": "Tu cédula:", "ejemplo": "V-18.901.234"},
            {"id": "APODERADO_NOMBRE",   "pregunta": "Nombre completo de la persona que recibirá el poder:", "ejemplo": "Pedro José Ramírez"},
            {"id": "APODERADO_CEDULA",   "pregunta": "Cédula del apoderado:", "ejemplo": "V-20.123.456"},
            {"id": "FACULTADES",         "pregunta": "Describe las facultades que le otorgas (qué puede hacer en tu nombre):", "ejemplo": "Realizar trámites bancarios, cobrar cheques, firmar documentos ante instituciones públicas y privadas, y representarme ante cualquier organismo del Estado"},
            {"id": "CIUDAD",             "pregunta": "Ciudad:", "ejemplo": "Maracaibo"},
            {"id": "FECHA_FIRMA",        "pregunta": "Fecha de hoy:", "ejemplo": "17 de marzo de 2026"},
        ]
    },
    "4": {
        "nombre": "Constancia de Residencia",
        "archivo": "constancia_residencia.docx",
        "descripcion": "Documento que certifica tu dirección de residencia",
        "campos": [
            {"id": "NOMBRE_COMPLETO",    "pregunta": "Tu nombre completo:", "ejemplo": "Luis Alberto Morales"},
            {"id": "CEDULA",             "pregunta": "Tu cédula:", "ejemplo": "V-22.345.678"},
            {"id": "DIRECCION",          "pregunta": "Tu dirección completa de residencia:", "ejemplo": "Calle 5, Casa N° 12, Sector El Paraíso"},
            {"id": "MUNICIPIO",          "pregunta": "Municipio:", "ejemplo": "Libertador"},
            {"id": "ESTADO",             "pregunta": "Estado:", "ejemplo": "Distrito Capital"},
            {"id": "TIEMPO_RESIDENCIA",  "pregunta": "Tiempo que llevas viviendo en esa dirección:", "ejemplo": "3 años"},
            {"id": "CIUDAD",             "pregunta": "Ciudad:", "ejemplo": "Caracas"},
            {"id": "FECHA_FIRMA",        "pregunta": "Fecha de hoy:", "ejemplo": "17 de marzo de 2026"},
        ]
    },
    "5": {
        "nombre": "Acta Constitutiva de Empresa (C.A.)",
        "archivo": "acta_constitutiva_ca.docx",
        "descripcion": "Documento para registrar una Compañía Anónima ante el Registro Mercantil",
        "campos": [
            {"id": "EMPRESA_NOMBRE",     "pregunta": "Nombre de la empresa (sin C.A.):", "ejemplo": "Tech Solutions"},
            {"id": "OBJETO_SOCIAL",      "pregunta": "Describe la actividad de la empresa (objeto social):", "ejemplo": "Consultoría en tecnología de información, desarrollo de software, diseño web, análisis de datos, y servicios de asesoría empresarial"},
            {"id": "SOCIO1_NOMBRE",      "pregunta": "Nombre completo del Socio 1:", "ejemplo": "Jonathan Lobo"},
            {"id": "SOCIO1_CEDULA",      "pregunta": "Cédula del Socio 1:", "ejemplo": "V-25.123.456"},
            {"id": "SOCIO1_PORCENTAJE",  "pregunta": "Porcentaje de participación del Socio 1:", "ejemplo": "50"},
            {"id": "SOCIO2_NOMBRE",      "pregunta": "Nombre completo del Socio 2:", "ejemplo": "María García"},
            {"id": "SOCIO2_CEDULA",      "pregunta": "Cédula del Socio 2:", "ejemplo": "V-26.789.012"},
            {"id": "SOCIO2_PORCENTAJE",  "pregunta": "Porcentaje de participación del Socio 2:", "ejemplo": "50"},
            {"id": "CAPITAL_SOCIAL",     "pregunta": "Capital social (monto):", "ejemplo": "1.000 dólares americanos"},
            {"id": "CAPITAL_LETRAS",     "pregunta": "Capital en letras:", "ejemplo": "un mil dólares americanos"},
            {"id": "DIRECCION_EMPRESA",  "pregunta": "Dirección de la empresa:", "ejemplo": "Centro Comercial Plaza, Nivel 2, Local 15"},
            {"id": "MUNICIPIO",          "pregunta": "Municipio:", "ejemplo": "Chacao"},
            {"id": "ESTADO",             "pregunta": "Estado:", "ejemplo": "Miranda"},
            {"id": "CIUDAD",             "pregunta": "Ciudad:", "ejemplo": "Caracas"},
            {"id": "FECHA_FIRMA",        "pregunta": "Fecha de hoy:", "ejemplo": "17 de marzo de 2026"},
        ]
    },
}


# ─── ESTADO DE CONVERSACIÓN POR USUARIO ───────────────────────────────────────
# Almacena en qué paso está cada usuario que está llenando un documento

conversaciones_doc = {}  # user_id -> {plantilla, campo_actual, datos}
menu_pendiente = set()   # user_ids que vieron el menú y deben elegir 1-5


def iniciar_documento(user_id: int, opcion: str) -> str | None:
    """Inicia el flujo de un documento. Retorna la primera pregunta o None si la opción es inválida."""
    if opcion not in PLANTILLAS:
        return None

    plantilla = PLANTILLAS[opcion]
    conversaciones_doc[user_id] = {
        "plantilla": opcion,
        "campo_actual": 0,
        "datos": {},
    }

    campo = plantilla["campos"][0]
    return (
        f"Vamos a preparar tu <b>{plantilla['nombre']}</b>.\n\n"
        f"Necesito {len(plantilla['campos'])} datos. Puedes escribir /cancelar en cualquier momento.\n\n"
        f"<b>{campo['pregunta']}</b>\n"
        f"<i>Ejemplo: {campo['ejemplo']}</i>"
    )


def esta_en_documento(user_id: int) -> bool:
    """Verifica si el usuario está en proceso de llenar un documento O eligiendo del menú."""
    return user_id in conversaciones_doc or user_id in menu_pendiente


def esperando_seleccion(user_id: int) -> bool:
    """Verifica si el usuario está eligiendo del menú."""
    return user_id in menu_pendiente


def marcar_menu_pendiente(user_id: int):
    """Marca que el usuario vio el menú y debe elegir."""
    menu_pendiente.add(user_id)


def procesar_seleccion_menu(user_id: int, texto: str) -> str | None:
    """Procesa la selección del menú. Retorna la primera pregunta o mensaje de error."""
    menu_pendiente.discard(user_id)
    opcion = texto.strip()
    if opcion in PLANTILLAS:
        return iniciar_documento(user_id, opcion)
    else:
        # Opción inválida — volver a mostrar menú
        menu_pendiente.add(user_id)
        opciones = ", ".join(PLANTILLAS.keys())
        return f"Opcion invalida. Escribe un numero valido ({opciones}) o /cancelar para salir."


def procesar_respuesta(user_id: int, texto: str) -> tuple[str, str | None]:
    """
    Procesa la respuesta del usuario al campo actual.
    Retorna: (mensaje_respuesta, ruta_archivo_o_None)
    - Si hay más campos: retorna la siguiente pregunta
    - Si terminó: genera el documento y retorna la ruta del archivo
    """
    if user_id not in conversaciones_doc:
        return ("No hay un documento en proceso. Usa /documento para empezar.", None)

    estado = conversaciones_doc[user_id]
    plantilla = PLANTILLAS[estado["plantilla"]]
    campo_actual = plantilla["campos"][estado["campo_actual"]]

    # Guardar el dato
    estado["datos"][campo_actual["id"]] = texto.strip()
    estado["campo_actual"] += 1

    # ¿Hay más campos?
    if estado["campo_actual"] < len(plantilla["campos"]):
        siguiente = plantilla["campos"][estado["campo_actual"]]
        progreso = estado["campo_actual"]
        total = len(plantilla["campos"])
        return (
            f"[{progreso}/{total}] <b>{siguiente['pregunta']}</b>\n"
            f"<i>Ejemplo: {siguiente['ejemplo']}</i>",
            None
        )

    # Todos los campos completos → generar documento
    try:
        ruta = generar_documento(plantilla, estado["datos"])
        del conversaciones_doc[user_id]
        return (
            f"Tu <b>{plantilla['nombre']}</b> esta listo.\n\n"
            "Revisalo con un abogado antes de firmar.",
            ruta
        )
    except Exception as e:
        logger.error(f"Error generando documento: {e}")
        del conversaciones_doc[user_id]
        return ("Hubo un error generando el documento. Intenta de nuevo con /documento.", None)


def cancelar_documento(user_id: int) -> str:
    """Cancela el documento en proceso."""
    menu_pendiente.discard(user_id)
    if user_id in conversaciones_doc:
        del conversaciones_doc[user_id]
        return "Documento cancelado. Puedes empezar otro con /documento."
    return "No hay un documento en proceso."


# ─── GENERADOR DE DOCUMENTOS ──────────────────────────────────────────────────

def generar_documento(plantilla: dict, datos: dict) -> str:
    """
    Abre la plantilla .docx, reemplaza los marcadores {{CAMPO}} con los datos
    del usuario, y guarda el documento generado.
    Retorna la ruta del archivo generado.
    """
    ruta_plantilla = os.path.join(PLANTILLAS_DIR, plantilla["archivo"])

    if not os.path.exists(ruta_plantilla):
        raise FileNotFoundError(f"Plantilla no encontrada: {ruta_plantilla}")

    doc = Document(ruta_plantilla)

    # Reemplazar en párrafos
    for parrafo in doc.paragraphs:
        _reemplazar_en_parrafo(parrafo, datos)

    # Reemplazar en tablas (si las hay)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    _reemplazar_en_parrafo(parrafo, datos)

    # Reemplazar en headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for parrafo in header.paragraphs:
                    _reemplazar_en_parrafo(parrafo, datos)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for parrafo in footer.paragraphs:
                    _reemplazar_en_parrafo(parrafo, datos)

    # Guardar en carpeta temporal
    output_dir = os.path.join(str(config.DATA_DIR), "documentos_generados")
    os.makedirs(output_dir, exist_ok=True)

    # Nombre del archivo: tipo_fecha_hora.docx
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_limpio = plantilla["archivo"].replace(".docx", "")
    nombre_output = f"{nombre_limpio}_{timestamp}.docx"
    ruta_output = os.path.join(output_dir, nombre_output)

    doc.save(ruta_output)
    logger.info(f"Documento generado: {ruta_output}")

    return ruta_output


def _reemplazar_en_parrafo(parrafo, datos: dict):
    """
    Reemplaza marcadores {{CAMPO}} en un párrafo, preservando el formato.
    Los marcadores pueden estar divididos entre varios runs, así que
    primero reconstruimos el texto completo y luego redistribuimos.
    """
    texto_completo = "".join(run.text for run in parrafo.runs)

    # ¿Hay algún marcador en este párrafo?
    if "{{" not in texto_completo:
        return

    # Reemplazar todos los marcadores
    texto_nuevo = texto_completo
    for campo_id, valor in datos.items():
        texto_nuevo = texto_nuevo.replace(f"{{{{{campo_id}}}}}", valor)

    # Si no cambió nada, no tocar
    if texto_nuevo == texto_completo:
        return

    # Redistribuir: poner todo el texto en el primer run, vaciar los demás
    if parrafo.runs:
        parrafo.runs[0].text = texto_nuevo
        for run in parrafo.runs[1:]:
            run.text = ""


def listar_plantillas() -> str:
    """Retorna el menú de plantillas disponibles formateado para Telegram."""
    lineas = []
    for key, p in PLANTILLAS.items():
        lineas.append(f"  {key}. <b>{p['nombre']}</b>\n     <i>{p['descripcion']}</i>")
    return "\n\n".join(lineas)
