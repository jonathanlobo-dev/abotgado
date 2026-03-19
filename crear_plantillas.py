"""
Genera las plantillas .docx para aBOTgado.
Ejecutar una sola vez: python crear_plantillas.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PLANTILLAS_DIR = os.path.join(os.path.dirname(__file__), "plantillas")
os.makedirs(PLANTILLAS_DIR, exist_ok=True)


def estilo_base(doc):
    """Configura estilos base del documento."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def titulo(doc, texto):
    """Agrega un título centrado y en negrita."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(18)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    return p


def subtitulo(doc, texto):
    """Agrega un subtítulo en negrita."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def parrafo(doc, texto):
    """Agrega un párrafo justificado."""
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def linea_firma(doc, nombre_placeholder, cedula_placeholder, rol):
    """Agrega bloque de firma."""
    doc.add_paragraph()  # espacio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_" * 40)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(f"{nombre_placeholder}")
    run.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"C.I.: {cedula_placeholder}")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(rol)
    p4.space_after = Pt(12)


def guardar(doc, nombre):
    ruta = os.path.join(PLANTILLAS_DIR, nombre)
    doc.save(ruta)
    print(f"  OK: {nombre}")


# ═══════════════════════════════════════════════════════════════════════════════
# 1. CONTRATO DE ARRENDAMIENTO DE VIVIENDA
# ═══════════════════════════════════════════════════════════════════════════════

def crear_contrato_arrendamiento():
    doc = Document()
    estilo_base(doc)

    titulo(doc, "CONTRATO DE ARRENDAMIENTO DE VIVIENDA")

    parrafo(doc,
        "Entre {{ARRENDADOR_NOMBRE}}, venezolano(a), mayor de edad, titular de la "
        "cédula de identidad N° {{ARRENDADOR_CEDULA}}, quien en lo adelante se "
        "denominará \"EL ARRENDADOR\"; y {{ARRENDATARIO_NOMBRE}}, venezolano(a), "
        "mayor de edad, titular de la cédula de identidad N° {{ARRENDATARIO_CEDULA}}, "
        "quien en lo adelante se denominará \"EL ARRENDATARIO\"; se ha convenido en "
        "celebrar el presente contrato de arrendamiento, el cual se regirá por las "
        "siguientes cláusulas:")

    subtitulo(doc, "CLÁUSULA PRIMERA: OBJETO DEL CONTRATO")
    parrafo(doc,
        "EL ARRENDADOR da en arrendamiento a EL ARRENDATARIO, quien lo recibe, un "
        "inmueble destinado a vivienda, ubicado en: {{DIRECCION_INMUEBLE}}, en la "
        "ciudad de {{CIUDAD}}, República Bolivariana de Venezuela.")

    subtitulo(doc, "CLÁUSULA SEGUNDA: DURACIÓN")
    parrafo(doc,
        "El presente contrato tendrá una duración de {{DURACION_MESES}} meses, "
        "contados a partir del {{FECHA_INICIO}}, pudiendo ser prorrogado por mutuo "
        "acuerdo de las partes, de conformidad con lo establecido en la Ley para la "
        "Regularización y Control de los Arrendamientos de Vivienda.")

    subtitulo(doc, "CLÁUSULA TERCERA: CANON DE ARRENDAMIENTO")
    parrafo(doc,
        "El canon de arrendamiento mensual se fija en la cantidad de {{CANON_MENSUAL}} "
        "({{CANON_LETRAS}}), pagaderos dentro de los primeros cinco (5) días de cada "
        "mes, mediante transferencia bancaria o depósito a la cuenta que indique "
        "EL ARRENDADOR.")

    subtitulo(doc, "CLÁUSULA CUARTA: DEPÓSITO EN GARANTÍA")
    parrafo(doc,
        "EL ARRENDATARIO entrega en este acto a EL ARRENDADOR la cantidad equivalente "
        "a un (1) mes de canon de arrendamiento como depósito en garantía, el cual será "
        "devuelto al término del contrato, previa verificación del estado del inmueble, "
        "de conformidad con el artículo 80 de la Ley para la Regularización y Control "
        "de los Arrendamientos de Vivienda.")

    subtitulo(doc, "CLÁUSULA QUINTA: USO DEL INMUEBLE")
    parrafo(doc,
        "EL ARRENDATARIO se compromete a destinar el inmueble exclusivamente para uso "
        "de vivienda, no pudiendo subarrendar total o parcialmente sin autorización "
        "escrita de EL ARRENDADOR.")

    subtitulo(doc, "CLÁUSULA SEXTA: CONSERVACIÓN DEL INMUEBLE")
    parrafo(doc,
        "EL ARRENDATARIO se obliga a mantener el inmueble en buen estado de "
        "conservación y a realizar las reparaciones menores que sean necesarias. Las "
        "reparaciones mayores serán por cuenta de EL ARRENDADOR, salvo que los daños "
        "hayan sido causados por negligencia de EL ARRENDATARIO.")

    subtitulo(doc, "CLÁUSULA SÉPTIMA: SERVICIOS PÚBLICOS")
    parrafo(doc,
        "EL ARRENDATARIO se obliga al pago puntual de los servicios públicos "
        "(electricidad, agua, aseo, gas, internet, teléfono y condominio si aplica) "
        "durante la vigencia del contrato.")

    subtitulo(doc, "CLÁUSULA OCTAVA: ENTREGA DEL INMUEBLE")
    parrafo(doc,
        "Al término del contrato, EL ARRENDATARIO se obliga a entregar el inmueble "
        "en las mismas condiciones en que lo recibió, salvo el desgaste natural por "
        "el uso normal, libre de personas y bienes que no formen parte del inmueble.")

    subtitulo(doc, "CLÁUSULA NOVENA: CAUSALES DE RESOLUCIÓN")
    parrafo(doc,
        "El presente contrato podrá resolverse por: a) Falta de pago de dos (2) "
        "cánones consecutivos; b) Uso del inmueble para fines distintos al convenido; "
        "c) Subarriendo no autorizado; d) Daños graves al inmueble; de conformidad "
        "con lo establecido en la ley vigente.")

    subtitulo(doc, "CLÁUSULA DÉCIMA: LEGISLACIÓN APLICABLE")
    parrafo(doc,
        "El presente contrato se rige por las disposiciones del Código Civil "
        "venezolano, la Ley para la Regularización y Control de los Arrendamientos "
        "de Vivienda, y la Ley contra el Desalojo Arbitrario de Viviendas. Para "
        "cualquier controversia, las partes se someten a los tribunales competentes "
        "de la ciudad de {{CIUDAD}}, previo agotamiento de la vía administrativa "
        "ante la SUNAVI.")

    parrafo(doc,
        "Se firman dos (2) ejemplares de un mismo tenor y a un solo efecto, en la "
        "ciudad de {{CIUDAD}}, a los {{FECHA_FIRMA}}.")

    linea_firma(doc, "{{ARRENDADOR_NOMBRE}}", "{{ARRENDADOR_CEDULA}}", "EL ARRENDADOR")
    linea_firma(doc, "{{ARRENDATARIO_NOMBRE}}", "{{ARRENDATARIO_CEDULA}}", "EL ARRENDATARIO")

    guardar(doc, "contrato_arrendamiento.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. CARTA DE RENUNCIA
# ═══════════════════════════════════════════════════════════════════════════════

def crear_carta_renuncia():
    doc = Document()
    estilo_base(doc)

    # Fecha y lugar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("{{CIUDAD}}, {{FECHA_FIRMA}}")

    doc.add_paragraph()  # espacio

    parrafo(doc, "Señores")
    p = doc.add_paragraph()
    run = p.add_run("{{EMPRESA_NOMBRE}}")
    run.bold = True
    parrafo(doc, "Presente.-")

    doc.add_paragraph()

    titulo(doc, "CARTA DE RENUNCIA VOLUNTARIA")

    parrafo(doc,
        "Yo, {{TRABAJADOR_NOMBRE}}, venezolano(a), mayor de edad, titular de la "
        "cédula de identidad N° {{TRABAJADOR_CEDULA}}, quien me desempeño en el "
        "cargo de {{CARGO}} en la empresa {{EMPRESA_NOMBRE}}, desde el "
        "{{FECHA_INGRESO}}, por medio de la presente manifiesto de manera formal "
        "y voluntaria mi decisión irrevocable de renunciar al cargo que vengo "
        "desempeñando.")

    parrafo(doc,
        "La presente renuncia tendrá efecto a partir del {{FECHA_RENUNCIA}}, "
        "de conformidad con lo establecido en el artículo 80 de la Ley Orgánica "
        "del Trabajo, los Trabajadores y las Trabajadoras (LOTTT).")

    parrafo(doc,
        "Solicito formalmente el cálculo y pago de mis prestaciones sociales y "
        "demás conceptos laborales que me corresponden por ley, incluyendo: "
        "prestaciones de antigüedad, vacaciones fraccionadas, bono vacacional "
        "fraccionado, utilidades fraccionadas, y cualquier otro concepto adeudado, "
        "dentro del plazo de cinco (5) días establecido en el artículo 142 de "
        "la LOTTT.")

    parrafo(doc,
        "Agradezco las oportunidades de crecimiento profesional brindadas durante "
        "el tiempo laborado en esta empresa.")

    parrafo(doc, "Sin otro particular, me despido atentamente.")

    linea_firma(doc, "{{TRABAJADOR_NOMBRE}}", "{{TRABAJADOR_CEDULA}}", "TRABAJADOR(A)")

    guardar(doc, "carta_renuncia.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. PODER NOTARIAL GENERAL
# ═══════════════════════════════════════════════════════════════════════════════

def crear_poder_notarial():
    doc = Document()
    estilo_base(doc)

    titulo(doc, "PODER GENERAL")

    parrafo(doc,
        "Yo, {{PODERDANTE_NOMBRE}}, venezolano(a), mayor de edad, de este domicilio, "
        "titular de la cédula de identidad N° {{PODERDANTE_CEDULA}}, actuando en mi "
        "propio nombre, por medio del presente documento declaro:")

    parrafo(doc,
        "Que confiero PODER GENERAL, amplio y suficiente cuanto en derecho se "
        "requiere, al ciudadano(a) {{APODERADO_NOMBRE}}, venezolano(a), mayor de "
        "edad, titular de la cédula de identidad N° {{APODERADO_CEDULA}}, para "
        "que en mi nombre y representación ejerza las siguientes facultades:")

    subtitulo(doc, "FACULTADES:")
    parrafo(doc, "{{FACULTADES}}")

    parrafo(doc,
        "El presente poder se otorga de conformidad con lo establecido en los "
        "artículos 1.684 y siguientes del Código Civil venezolano, y podrá ser "
        "revocado en cualquier momento mediante notificación formal al apoderado.")

    parrafo(doc,
        "El apoderado queda facultado para sustituir total o parcialmente el "
        "presente poder en la persona o personas que considere conveniente, "
        "reservándose siempre el ejercicio del mismo.")

    parrafo(doc,
        "Es en la ciudad de {{CIUDAD}}, a los {{FECHA_FIRMA}}.")

    linea_firma(doc, "{{PODERDANTE_NOMBRE}}", "{{PODERDANTE_CEDULA}}", "EL PODERDANTE")
    linea_firma(doc, "{{APODERADO_NOMBRE}}", "{{APODERADO_CEDULA}}", "EL APODERADO (Acepta)")

    # Espacio para notaría
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTA DE AUTENTICACIÓN")
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("(Espacio reservado para la Notaría Pública)")

    guardar(doc, "poder_notarial.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 4. CONSTANCIA DE RESIDENCIA
# ═══════════════════════════════════════════════════════════════════════════════

def crear_constancia_residencia():
    doc = Document()
    estilo_base(doc)

    titulo(doc, "CONSTANCIA DE RESIDENCIA")

    parrafo(doc,
        "Quien suscribe, hace constar por medio de la presente que el(la) "
        "ciudadano(a) {{NOMBRE_COMPLETO}}, venezolano(a), mayor de edad, titular "
        "de la cédula de identidad N° {{CEDULA}}, reside en la siguiente dirección:")

    # Dirección destacada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    p.space_after = Pt(12)
    run = p.add_run("{{DIRECCION}}")
    run.bold = True
    run.font.size = Pt(12)

    parrafo(doc,
        "Municipio {{MUNICIPIO}}, Estado {{ESTADO}}, República Bolivariana de "
        "Venezuela.")

    parrafo(doc,
        "El(la) mencionado(a) ciudadano(a) tiene un tiempo de residencia en "
        "dicha dirección de {{TIEMPO_RESIDENCIA}}.")

    parrafo(doc,
        "Constancia que se expide a solicitud de la parte interesada, en la "
        "ciudad de {{CIUDAD}}, a los {{FECHA_FIRMA}}.")

    doc.add_paragraph()
    doc.add_paragraph()

    linea_firma(doc, "{{NOMBRE_COMPLETO}}", "{{CEDULA}}", "SOLICITANTE")

    # Espacio para testigos
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TESTIGOS:")
    run.bold = True

    for i in range(1, 3):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_" * 40)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(f"Testigo {i} — Nombre y C.I.")

    guardar(doc, "constancia_residencia.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 5. ACTA CONSTITUTIVA DE COMPAÑÍA ANÓNIMA (C.A.)
# ═══════════════════════════════════════════════════════════════════════════════

def crear_acta_constitutiva():
    doc = Document()
    estilo_base(doc)

    titulo(doc, "ACTA CONSTITUTIVA Y ESTATUTOS SOCIALES")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(18)
    run = p.add_run("{{EMPRESA_NOMBRE}}, C.A.")
    run.bold = True
    run.font.size = Pt(16)

    parrafo(doc,
        "Nosotros, {{SOCIO1_NOMBRE}}, venezolano(a), mayor de edad, titular de la "
        "cédula de identidad N° {{SOCIO1_CEDULA}}; y {{SOCIO2_NOMBRE}}, "
        "venezolano(a), mayor de edad, titular de la cédula de identidad "
        "N° {{SOCIO2_CEDULA}}; hemos convenido en constituir, como en efecto "
        "constituimos por medio del presente documento, una compañía anónima que "
        "se regirá por las disposiciones del Código de Comercio y por las "
        "cláusulas siguientes:")

    # --- TÍTULO I ---
    subtitulo(doc, "TÍTULO I — DENOMINACIÓN, DOMICILIO, OBJETO Y DURACIÓN")

    subtitulo(doc, "CLÁUSULA PRIMERA: DENOMINACIÓN")
    parrafo(doc,
        "La compañía se denominará \"{{EMPRESA_NOMBRE}}, C.A.\", y podrá usar "
        "comercialmente la abreviatura de su nombre, de conformidad con el "
        "artículo 202 del Código de Comercio.")

    subtitulo(doc, "CLÁUSULA SEGUNDA: DOMICILIO")
    parrafo(doc,
        "El domicilio principal de la compañía será en: {{DIRECCION_EMPRESA}}, "
        "Municipio {{MUNICIPIO}}, Estado {{ESTADO}}, República Bolivariana de "
        "Venezuela, pudiendo establecer sucursales, agencias u oficinas en "
        "cualquier lugar del país o del exterior.")

    subtitulo(doc, "CLÁUSULA TERCERA: OBJETO SOCIAL")
    parrafo(doc,
        "La compañía tendrá por objeto: {{OBJETO_SOCIAL}}. Además, podrá "
        "realizar cualquier otra actividad de lícito comercio conexa o no con "
        "su objeto principal.")

    subtitulo(doc, "CLÁUSULA CUARTA: DURACIÓN")
    parrafo(doc,
        "La duración de la compañía será de cincuenta (50) años, contados a "
        "partir de su inscripción en el Registro Mercantil, pudiendo prorrogarse "
        "por decisión de la Asamblea de Accionistas.")

    # --- TÍTULO II ---
    subtitulo(doc, "TÍTULO II — CAPITAL SOCIAL Y ACCIONES")

    subtitulo(doc, "CLÁUSULA QUINTA: CAPITAL SOCIAL")
    parrafo(doc,
        "El capital social de la compañía es de {{CAPITAL_SOCIAL}} "
        "({{CAPITAL_LETRAS}}), dividido en acciones nominativas de igual valor, "
        "el cual ha sido suscrito y pagado en su totalidad de la siguiente manera:")

    parrafo(doc,
        "• {{SOCIO1_NOMBRE}} (C.I. {{SOCIO1_CEDULA}}): suscribe y paga el "
        "{{SOCIO1_PORCENTAJE}}% del capital social.")
    parrafo(doc,
        "• {{SOCIO2_NOMBRE}} (C.I. {{SOCIO2_CEDULA}}): suscribe y paga el "
        "{{SOCIO2_PORCENTAJE}}% del capital social.")

    subtitulo(doc, "CLÁUSULA SEXTA: ACCIONES")
    parrafo(doc,
        "Las acciones serán nominativas y no podrán ser cedidas ni traspasadas "
        "sin el consentimiento previo de la Asamblea de Accionistas, quienes "
        "tendrán derecho preferente de adquisición.")

    # --- TÍTULO III ---
    subtitulo(doc, "TÍTULO III — ADMINISTRACIÓN Y DIRECCIÓN")

    subtitulo(doc, "CLÁUSULA SÉPTIMA: DIRECCIÓN")
    parrafo(doc,
        "La dirección y administración de la compañía estará a cargo de una "
        "Junta Directiva compuesta por un(a) Director(a) General y un(a) "
        "Director(a) Ejecutivo(a), quienes serán designados por la Asamblea de "
        "Accionistas y durarán cinco (5) años en sus funciones, pudiendo ser "
        "reelegidos.")

    subtitulo(doc, "CLÁUSULA OCTAVA: REPRESENTACIÓN")
    parrafo(doc,
        "El(la) Director(a) General tendrá la representación legal de la compañía "
        "y podrá ejecutar todos los actos de administración y disposición, "
        "incluyendo: abrir y cerrar cuentas bancarias, suscribir contratos, "
        "otorgar poderes, comparecer en juicio, y en general, realizar todo acto "
        "necesario para el cumplimiento del objeto social.")

    # --- TÍTULO IV ---
    subtitulo(doc, "TÍTULO IV — ASAMBLEA DE ACCIONISTAS")

    subtitulo(doc, "CLÁUSULA NOVENA: ASAMBLEA")
    parrafo(doc,
        "La Asamblea General de Accionistas es el órgano supremo de la compañía. "
        "Las asambleas ordinarias se celebrarán dentro de los noventa (90) días "
        "siguientes al cierre del ejercicio económico anual. Las asambleas "
        "extraordinarias se celebrarán cuando así lo requieran los intereses "
        "de la compañía.")

    # --- TÍTULO V ---
    subtitulo(doc, "TÍTULO V — EJERCICIO ECONÓMICO Y UTILIDADES")

    subtitulo(doc, "CLÁUSULA DÉCIMA: EJERCICIO ECONÓMICO")
    parrafo(doc,
        "El ejercicio económico de la compañía comenzará el primero (1°) de "
        "enero y terminará el treinta y uno (31) de diciembre de cada año.")

    subtitulo(doc, "CLÁUSULA DÉCIMA PRIMERA: COMISARIO")
    parrafo(doc,
        "La compañía tendrá un(a) Comisario(a), quien deberá ser Contador(a) "
        "Público(a) Colegiado(a), designado(a) por la Asamblea de Accionistas, "
        "de conformidad con los artículos 287 y siguientes del Código de Comercio.")

    # --- DESIGNACIONES ---
    subtitulo(doc, "DESIGNACIONES")
    parrafo(doc,
        "Los socios acuerdan designar como Director(a) General a {{SOCIO1_NOMBRE}}, "
        "y como Director(a) Ejecutivo(a) a {{SOCIO2_NOMBRE}}, quienes declaran "
        "aceptar las designaciones y se comprometen a cumplir con las obligaciones "
        "inherentes a sus cargos.")

    parrafo(doc,
        "Se firma en la ciudad de {{CIUDAD}}, a los {{FECHA_FIRMA}}.")

    linea_firma(doc, "{{SOCIO1_NOMBRE}}", "{{SOCIO1_CEDULA}}", "SOCIO / DIRECTOR(A) GENERAL")
    linea_firma(doc, "{{SOCIO2_NOMBRE}}", "{{SOCIO2_CEDULA}}", "SOCIO / DIRECTOR(A) EJECUTIVO(A)")

    # Nota para Registro Mercantil
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTA DE REGISTRO MERCANTIL")
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("(Espacio reservado para el Registro Mercantil)")

    guardar(doc, "acta_constitutiva_ca.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Creando plantillas de documentos legales...\n")

    crear_contrato_arrendamiento()
    crear_carta_renuncia()
    crear_poder_notarial()
    crear_constancia_residencia()
    crear_acta_constitutiva()

    print(f"\nListo! 5 plantillas creadas en: {PLANTILLAS_DIR}")
    print("   Ahora puedes revisarlas y ajustarlas con tu abogado.")
